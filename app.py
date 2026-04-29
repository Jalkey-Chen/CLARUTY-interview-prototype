"""Streamlit entrypoint for the CLARITY Patient Explainer Dashboard prototype."""

import json
import os
import zipfile
import urllib.error
import urllib.request
from hashlib import sha256
from io import BytesIO
from html import escape
from pathlib import Path
from xml.etree import ElementTree

import streamlit as st
from docx import Document
from dotenv import load_dotenv


APP_ROOT = Path(__file__).resolve().parent
DATA_DIR = APP_ROOT / "data"
SAMPLE_NOTE_PATH = DATA_DIR / "sample_case_note.txt"
FACT_BASE_PATH = DATA_DIR / "extracted_fact_base.json"
VERSION_METADATA_PATH = DATA_DIR / "version_metadata.json"
WRITEUP_DIR = APP_ROOT / "writeup"
LIMITATIONS_PATH = WRITEUP_DIR / "limitations_and_next_steps.md"
ENV_PATH = APP_ROOT / ".env"
API_TIMEOUT_SECONDS = 60
API_ENV_VARS = {
    "OPENAI_API_KEY": "LLM/script generation provider key",
    "NOTE_INGESTION_API_URL": "Future Step 1 secure note ingestion endpoint",
    "FACT_EXTRACTION_API_URL": "Future Step 2 fact extraction endpoint",
    "SCRIPT_GENERATION_API_URL": "Future Step 3 script generation endpoint",
    "VERIFICATION_API_URL": "Future script verification endpoint",
    "VIDEO_GENERATION_API_URL": "Future Step 4 video generation endpoint",
}


def inject_custom_css() -> None:
    """Apply lightweight visual styling for the Streamlit prototype.

    Args:
        None.

    Returns:
        None. CSS is injected into the Streamlit page as a rendering side
        effect.

    CLARITY pipeline role:
        Improves the interview demo experience without changing the underlying
        prototype data flow. The CSS creates clearer hierarchy, workflow
        markers, compact fact cards, and a more intentional placeholder style so
        evaluators can understand the CLARITY flow quickly.
    """
    st.markdown(
        """
<style>
    .block-container {
        max-width: 1180px;
        padding-top: 1.6rem;
        padding-bottom: 3rem;
    }

    h1 {
        color: #12343b;
        letter-spacing: 0;
        margin-bottom: 0.2rem;
    }

    h2, h3 {
        color: #16323b;
        letter-spacing: 0;
    }

    .clarity-hero {
        border: 1px solid #d9e2e7;
        border-left: 6px solid #2b7a78;
        background: linear-gradient(135deg, #f7fbfb 0%, #ffffff 62%);
        padding: 1.1rem 1.25rem;
        border-radius: 8px;
        margin: 1rem 0 1.1rem 0;
    }

    .clarity-hero-title {
        color: #12343b;
        font-size: 1.05rem;
        font-weight: 700;
        margin-bottom: 0.25rem;
    }

    .clarity-hero-text {
        color: #38545c;
        font-size: 0.95rem;
        line-height: 1.45;
        margin: 0;
    }

    .workflow-strip {
        display: grid;
        grid-template-columns: repeat(5, minmax(0, 1fr));
        gap: 0.55rem;
        margin: 0.5rem 0 1.35rem 0;
    }

    .workflow-step {
        border: 1px solid #dbe5e8;
        background: #fbfcfd;
        border-radius: 8px;
        padding: 0.7rem 0.75rem;
        min-height: 76px;
    }

    .workflow-step-number {
        color: #2b7a78;
        font-size: 0.75rem;
        font-weight: 700;
        text-transform: uppercase;
        margin-bottom: 0.25rem;
    }

    .workflow-step-title {
        color: #16323b;
        font-size: 0.88rem;
        font-weight: 650;
        line-height: 1.25;
    }

    .step-heading {
        border-bottom: 1px solid #d9e2e7;
        margin: 0.15rem 0 1rem 0;
        padding-bottom: 0.65rem;
    }

    .step-kicker {
        color: #2b7a78;
        font-size: 0.78rem;
        font-weight: 750;
        text-transform: uppercase;
        margin-bottom: 0.15rem;
    }

    .step-title {
        color: #12343b;
        font-size: 1.28rem;
        font-weight: 750;
        margin-bottom: 0.2rem;
    }

    .step-summary {
        color: #526b73;
        font-size: 0.94rem;
        line-height: 1.45;
        margin: 0;
    }

    .fact-card {
        border: 1px solid #d9e2e7;
        background: #ffffff;
        border-radius: 8px;
        padding: 0.85rem 0.9rem;
        min-height: 96px;
        margin-bottom: 0.7rem;
    }

    .fact-label {
        color: #60777f;
        font-size: 0.72rem;
        font-weight: 750;
        text-transform: uppercase;
        margin-bottom: 0.25rem;
    }

    .fact-value {
        color: #17333b;
        font-size: 0.96rem;
        font-weight: 600;
        line-height: 1.35;
    }

    .compact-list ul {
        margin-top: 0.25rem;
        padding-left: 1.05rem;
    }

    .compact-list li {
        margin-bottom: 0.3rem;
    }

    .muted-note {
        color: #526b73;
        font-size: 0.88rem;
        line-height: 1.4;
    }

    .mode-summary {
        border: 1px solid #d9e2e7;
        background: #ffffff;
        border-radius: 8px;
        padding: 1rem;
        margin-bottom: 0.75rem;
    }

    .mode-title {
        color: #12343b;
        font-size: 1.15rem;
        font-weight: 750;
        margin-bottom: 0.35rem;
    }

    .video-placeholder {
        border: 1px dashed #8ab7b5;
        background: #f6fbfb;
        border-radius: 8px;
        padding: 1.15rem;
        min-height: 180px;
        display: flex;
        flex-direction: column;
        justify-content: center;
    }

    .video-placeholder-title {
        color: #12343b;
        font-weight: 750;
        font-size: 1rem;
        margin-bottom: 0.4rem;
    }

    .video-placeholder-text {
        color: #526b73;
        line-height: 1.45;
        margin: 0;
    }

    @media (max-width: 900px) {
        .workflow-strip {
            grid-template-columns: 1fr;
        }
    }
</style>
""",
        unsafe_allow_html=True,
    )


def load_environment_config(env_path: Path) -> dict:
    """Load local environment variables for future API-backed steps.

    Args:
        env_path: Path to the local `.env` file. The file is intentionally not
        committed; `.env.example` documents the expected variable names.

    Returns:
        A dictionary containing configured values for the API-related variables
        used by future CLARITY integrations. Missing variables are represented
        as empty strings.

    CLARITY pipeline role:
        Prepares the prototype for Step 1-4 API integration without calling any
        external services in the current demo. This gives evaluators and future
        developers a clear place to configure ingestion, extraction, script
        generation, verification, and video generation services.
    """
    if env_path.exists():
        load_dotenv(env_path, override=True)

    return {name: os.getenv(name, "") for name in API_ENV_VARS}


def display_api_config_status(api_config: dict) -> None:
    """Display non-secret API configuration readiness in the sidebar.

    Args:
        api_config: Mapping returned by `load_environment_config`.

    Returns:
        None. The status is rendered into the Streamlit sidebar.

    CLARITY pipeline role:
        Makes it explicit that the current app is API-ready but still operates
        with local cached assets. Values are never printed, only whether each
        expected variable is configured.
    """
    with st.sidebar.expander("API configuration", expanded=False):
        st.caption(
            "Optional for this prototype. Add values to `.env` using "
            "`.env.example` as the template when connecting Steps 1-4 to APIs."
        )
        for variable_name, description in API_ENV_VARS.items():
            configured = bool(api_config.get(variable_name))
            status_label = "Configured" if configured else "Not set"
            status_icon = "OK" if configured else "Missing"
            st.markdown(f"**{variable_name}**: {status_icon} - {status_label}")
            st.caption(description)


def get_note_hash(note_text: str) -> str:
    """Create a stable short hash for a loaded clinical note.

    Args:
        note_text: Raw imported clinical note text.

    Returns:
        A short SHA-256 digest that can be used as a session-state cache key.

    CLARITY pipeline role:
        Lets the API-backed prototype avoid repeating the same extraction or
        generation calls on every Streamlit rerun for the same uploaded note.
    """
    return sha256(note_text.encode("utf-8")).hexdigest()[:16]


def post_json_to_api(endpoint_url: str, payload: dict, api_key: str = "") -> dict:
    """POST JSON to an API endpoint and parse a JSON object response.

    Args:
        endpoint_url: Fully qualified API endpoint URL.
        payload: JSON-serializable request body.
        api_key: Optional bearer token used for future hosted API providers.

    Returns:
        Parsed JSON object response. Returns an empty dictionary when the
        request fails, the response is not valid JSON, or the response JSON is
        not an object.

    CLARITY pipeline role:
        Provides a small API bridge for uploaded-note mode. The local demo uses
        cached files, but uploaded cases should move through API-backed Step 2
        fact extraction, Step 3 script generation/verification, and Step 4 video
        generation or retrieval.
    """
    request_body = json.dumps(payload).encode("utf-8")
    headers = {
        "Content-Type": "application/json",
        "Accept": "application/json",
    }
    if api_key:
        headers["Authorization"] = f"Bearer {api_key}"

    request = urllib.request.Request(
        endpoint_url,
        data=request_body,
        headers=headers,
        method="POST",
    )

    try:
        with urllib.request.urlopen(request, timeout=API_TIMEOUT_SECONDS) as response:
            response_text = response.read().decode("utf-8")
    except urllib.error.HTTPError as exc:
        error_text = exc.read().decode("utf-8", errors="replace")
        st.error(f"API request failed with HTTP {exc.code}: {error_text}")
        return {}
    except urllib.error.URLError as exc:
        st.error(f"API request failed: {exc.reason}")
        return {}
    except TimeoutError:
        st.error("API request timed out.")
        return {}

    try:
        data = json.loads(response_text)
    except json.JSONDecodeError as exc:
        st.error(f"API response was not valid JSON: {exc}")
        return {}

    if not isinstance(data, dict):
        st.error("API response must be a JSON object.")
        return {}
    return data


def extract_response_object(response: dict, keys: list[str]) -> dict:
    """Extract a nested dictionary from an API response.

    Args:
        response: Parsed API response object.
        keys: Candidate keys that may contain the desired nested object.

    Returns:
        The first nested dictionary found under the candidate keys, or the
        original response if none of those keys are present.

    CLARITY pipeline role:
        Makes the prototype tolerant of common API response envelopes such as
        `{ "fact_base": {...} }` while still accepting direct JSON objects.
    """
    for key in keys:
        value = response.get(key)
        if isinstance(value, dict):
            return value
    return response


def extract_response_text(response: dict, keys: list[str]) -> str:
    """Extract text content from an API response.

    Args:
        response: Parsed API response object.
        keys: Candidate string keys, ordered by preference.

    Returns:
        The first non-empty string found under the candidate keys, or an empty
        string if the response does not include generated text.

    CLARITY pipeline role:
        Lets the Step 4 script panel accept several likely API response shapes
        without forcing the first backend prototype to use one exact field name.
    """
    for key in keys:
        value = response.get(key)
        if isinstance(value, str) and value.strip():
            return value.strip()
    return ""


def call_fact_extraction_api(note_text: str, api_config: dict) -> dict:
    """Call the configured fact extraction API for an uploaded case.

    Args:
        note_text: Raw text extracted from the uploaded clinical note.
        api_config: Environment-derived API configuration.

    Returns:
        Structured fact base dictionary returned by the API, or an empty
        dictionary when the endpoint is missing or the request fails.

    CLARITY pipeline role:
        Implements API mode for Step 2. Uploaded notes should not reuse the
        sample fact base; they need a case-specific fact extraction response.
    """
    endpoint_url = api_config.get("FACT_EXTRACTION_API_URL", "")
    if not endpoint_url:
        st.warning(
            "API mode is active, but `FACT_EXTRACTION_API_URL` is not configured."
        )
        return {}

    response = post_json_to_api(
        endpoint_url,
        {"clinical_note": note_text},
        api_config.get("OPENAI_API_KEY", ""),
    )
    return extract_response_object(response, ["fact_base", "data", "result"])


def call_note_ingestion_api(note_text: str, api_config: dict) -> dict:
    """Call the optional note ingestion API for an uploaded case.

    Args:
        note_text: Raw text extracted from the uploaded clinical note.
        api_config: Environment-derived API configuration.

    Returns:
        Ingestion API response dictionary, or an empty dictionary when the
        endpoint is missing or the request fails.

    CLARITY pipeline role:
        Represents API mode for Step 1. The prototype can still pass extracted
        note text directly to downstream APIs, but this hook documents and
        exercises the intended secure ingestion boundary for uploaded cases.
    """
    endpoint_url = api_config.get("NOTE_INGESTION_API_URL", "")
    if not endpoint_url:
        st.warning(
            "API mode is active, but `NOTE_INGESTION_API_URL` is not configured; "
            "using local uploaded text for downstream API payloads."
        )
        return {}

    return post_json_to_api(
        endpoint_url,
        {"clinical_note": note_text},
        api_config.get("OPENAI_API_KEY", ""),
    )


def call_script_generation_api(
    note_text: str,
    fact_base: dict,
    mode_key: str,
    mode_metadata: dict,
    preferences: dict,
    api_config: dict,
) -> str:
    """Call the configured script generation API for an uploaded case.

    Args:
        note_text: Raw imported clinical note text.
        fact_base: Structured fact base from Step 2.
        mode_key: Selected explanation mode identifier.
        mode_metadata: Selected explanation mode metadata.
        preferences: Placeholder personalization selections from the sidebar.
        api_config: Environment-derived API configuration.

    Returns:
        Generated script markdown/text, or an empty string when unavailable.

    CLARITY pipeline role:
        Implements API mode for the script portion of Step 4. This keeps
        uploaded cases from showing static demo scripts that belong to the
        sample case.
    """
    endpoint_url = api_config.get("SCRIPT_GENERATION_API_URL", "")
    if not endpoint_url:
        st.warning(
            "API mode is active, but `SCRIPT_GENERATION_API_URL` is not configured."
        )
        return ""

    response = post_json_to_api(
        endpoint_url,
        {
            "clinical_note": note_text,
            "fact_base": fact_base,
            "mode_key": mode_key,
            "mode_metadata": mode_metadata,
            "preferences": preferences,
        },
        api_config.get("OPENAI_API_KEY", ""),
    )
    return extract_response_text(
        response,
        ["script_markdown", "script", "transcript", "content", "text"],
    )


def call_verification_api(
    note_text: str,
    fact_base: dict,
    script_text: str,
    mode_key: str,
    api_config: dict,
) -> dict:
    """Call the optional verification API for an uploaded-case script.

    Args:
        note_text: Raw imported clinical note text.
        fact_base: Structured fact base from Step 2.
        script_text: Script generated for the selected mode.
        mode_key: Selected explanation mode identifier.
        api_config: Environment-derived API configuration.

    Returns:
        Verification report dictionary, or an empty dictionary when the endpoint
        is not configured or the request fails.

    CLARITY pipeline role:
        Supports the intended CLARITY safeguard between script generation and
        video generation. The current UI displays the returned report but does
        not block video requests unless future API logic chooses to do so.
    """
    endpoint_url = api_config.get("VERIFICATION_API_URL", "")
    if not endpoint_url:
        st.warning(
            "API mode is active, but `VERIFICATION_API_URL` is not configured; "
            "script verification was skipped."
        )
        return {}

    return post_json_to_api(
        endpoint_url,
        {
            "clinical_note": note_text,
            "fact_base": fact_base,
            "script": script_text,
            "mode_key": mode_key,
        },
        api_config.get("OPENAI_API_KEY", ""),
    )


def call_video_generation_api(
    fact_base: dict,
    script_text: str,
    mode_key: str,
    mode_metadata: dict,
    verification_report: dict,
    api_config: dict,
) -> dict:
    """Call the configured video generation or retrieval API.

    Args:
        fact_base: Structured fact base from Step 2.
        script_text: Generated script text for the selected mode.
        mode_key: Selected explanation mode identifier.
        mode_metadata: Selected explanation mode metadata.
        verification_report: Optional verification report from the verification
        API.
        api_config: Environment-derived API configuration.

    Returns:
        Video API response dictionary. Common supported fields include
        `video_url`, `video_file`, `video_path`, `status`, and `message`.

    CLARITY pipeline role:
        Implements API mode for the video portion of Step 4. Uploaded cases
        should request or retrieve video outputs from the configured service
        instead of showing sample cached videos.
    """
    endpoint_url = api_config.get("VIDEO_GENERATION_API_URL", "")
    if not endpoint_url:
        st.warning(
            "API mode is active, but `VIDEO_GENERATION_API_URL` is not configured."
        )
        return {}

    return post_json_to_api(
        endpoint_url,
        {
            "fact_base": fact_base,
            "script": script_text,
            "mode_key": mode_key,
            "mode_metadata": mode_metadata,
            "verification_report": verification_report,
        },
        api_config.get("OPENAI_API_KEY", ""),
    )


def render_workflow_strip() -> None:
    """Render the five-step CLARITY workflow as a compact visual map.

    Args:
        None.

    Returns:
        None. The workflow strip is rendered directly into the page.

    CLARITY pipeline role:
        Gives evaluators a clear mental model of the prototype before they
        interact with any controls. This makes the app feel like a guided
        workflow rather than a loose collection of dashboard panels.
    """
    steps = [
        "Import Clinical Note",
        "Review Shared Fact Base",
        "Choose Explanation Mode",
        "Review Script and Cached Video",
        "Transparency and Next Steps",
    ]
    step_html = "".join(
        (
            '<div class="workflow-step">'
            f'<div class="workflow-step-number">Step {index}</div>'
            f'<div class="workflow-step-title">{escape(title)}</div>'
            "</div>"
        )
        for index, title in enumerate(steps, start=1)
    )
    st.markdown(
        f'<div class="workflow-strip">{step_html}</div>',
        unsafe_allow_html=True,
    )


def render_step_header(step_number: int, title: str, summary: str) -> None:
    """Render a consistent heading for each workflow step.

    Args:
        step_number: Numeric position in the CLARITY prototype workflow.
        title: Short step title displayed to the evaluator.
        summary: One-sentence explanation of what the step demonstrates.

    Returns:
        None. The heading is rendered directly into the page.

    CLARITY pipeline role:
        Creates consistent navigation cues across the dashboard so the demo
        reads as an end-to-end patient explainer pipeline.
    """
    st.markdown(
        f"""
        <div class="step-heading">
            <div class="step-kicker">Step {step_number}</div>
            <div class="step-title">{escape(title)}</div>
            <p class="step-summary">{escape(summary)}</p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_fact_card(label: str, value: str) -> None:
    """Render a compact labeled fact card.

    Args:
        label: The short label for the fact.
        value: The fact value to display.

    Returns:
        None. The card is rendered directly into the page.

    CLARITY pipeline role:
        Makes the shared fact base easier to scan by turning dense clinical
        fields into consistent, patient-facing units.
    """
    st.markdown(
        f"""
        <div class="fact-card">
            <div class="fact-label">{escape(label)}</div>
            <div class="fact-value">{escape(value)}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def load_json(path: Path) -> dict:
    """Load a JSON object from disk with Streamlit-friendly error handling.

    Args:
        path: Path to a JSON file expected to contain a top-level object.

    Returns:
        A dictionary parsed from the JSON file. Returns an empty dictionary when
        the file is missing, unreadable, malformed, or does not contain a JSON
        object.

    CLARITY pipeline role:
        Provides a safe loader for structured prototype assets such as the
        clinician-verifiable fact base and, in later milestones, version
        metadata. Safe loading matters because this dashboard is meant to stay
        demoable even while scripts, prompts, and videos are being iterated.
    """
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError:
        # A missing data file should be visible to the evaluator without
        # collapsing the rest of the prototype page.
        st.warning(f"Required JSON file was not found: {path}")
        return {}
    except json.JSONDecodeError as exc:
        st.error(f"Invalid JSON in {path.name}: {exc}")
        return {}
    except OSError as exc:
        st.error(f"Unable to read JSON file {path.name}: {exc}")
        return {}

    if not isinstance(data, dict):
        st.error(f"Expected {path.name} to contain a JSON object.")
        return {}
    return data


def load_markdown(path: Path) -> str:
    """Load a markdown document from disk with graceful error handling.

    Args:
        path: Path to the markdown file that should be rendered in the app.

    Returns:
        The markdown text when it can be read, or an empty string when the file
        is missing or unreadable.

    CLARITY pipeline role:
        Supports transcript, prompt, and writeup display throughout the
        prototype. Markdown assets are intentionally stored as editable files so
        scripts and documentation can be revised without changing the Streamlit
        application logic.
    """
    try:
        return path.read_text(encoding="utf-8")
    except FileNotFoundError:
        # Missing files are expected during iteration; the dashboard should
        # point to the missing asset instead of failing mid-demo.
        st.warning(f"Markdown file was not found: {path}")
    except OSError as exc:
        st.error(f"Unable to read markdown file {path.name}: {exc}")
    return ""


def load_sample_note(path: Path) -> str:
    """Load the built-in de-identified sample clinical note.

    Args:
        path: Path to the sample clinical note text file.

    Returns:
        The note text when the file can be read, or an empty string when the
        file is missing or unreadable.

    CLARITY pipeline role:
        Supports Step 1 of the prototype workflow by giving interview
        evaluators a reliable built-in case to load before they test upload
        behavior. This keeps the demo runnable even when the evaluator does not
        have a local clinical note prepared.
    """
    try:
        return path.read_text(encoding="utf-8")
    except FileNotFoundError:
        # Missing demo assets should not crash the app; graceful fallback keeps
        # the prototype usable while clearly showing what needs to be restored.
        st.warning(f"Sample note file was not found: {path}")
    except OSError as exc:
        st.error(f"Unable to read sample note: {exc}")
    return ""


def extract_docx_text_from_xml(file_bytes: bytes) -> str:
    """Extract Word text nodes directly from the `.docx` XML package.

    Args:
        file_bytes: Raw bytes from the uploaded Word document.

    Returns:
        Plain text assembled from WordprocessingML text nodes across document
        body, headers, footers, footnotes, endnotes, and other Word XML parts.
        Returns an empty string when no text nodes are present.

    CLARITY pipeline role:
        Provides a fallback parser for Step 1 when clinical note text is stored
        in Word structures that `python-docx` does not expose through ordinary
        paragraphs or tables. This keeps the prototype more robust for real
        interview demo files while still normalizing everything to plain text
        before API-backed fact extraction.
    """
    text_blocks: list[str] = []
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    with zipfile.ZipFile(BytesIO(file_bytes)) as archive:
        xml_part_names = [
            name
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        ]

        for part_name in xml_part_names:
            try:
                root = ElementTree.fromstring(archive.read(part_name))
            except ElementTree.ParseError:
                continue

            text_values = [
                node.text
                for node in root.findall(".//w:t", namespace)
                if node.text and node.text.strip()
            ]
            if text_values:
                text_blocks.append(" ".join(value.strip() for value in text_values))

    return "\n\n".join(text_blocks)


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract readable text from an uploaded Word `.docx` document.

    Args:
        file_bytes: Raw bytes from the uploaded Word document.

    Returns:
        Plain text extracted from document paragraphs, table cells, and a
        WordprocessingML XML fallback. Returns an empty string when the document
        has no readable text nodes.

    CLARITY pipeline role:
        Extends Step 1 so clinical notes can enter the prototype as Word
        documents, which is a more realistic format for interview demos and
        clinical handoffs. The function still converts the upload into plain
        text before later steps, keeping fact extraction separate from document
        format parsing.
    """
    document = Document(BytesIO(file_bytes))
    text_blocks: list[str] = []

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()
        if paragraph_text:
            text_blocks.append(paragraph_text)

    for table in document.tables:
        for row in table.rows:
            cell_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if cell_text:
                text_blocks.append(cell_text)

    paragraph_table_text = "\n\n".join(text_blocks)
    if paragraph_table_text:
        return paragraph_table_text

    # Some Word files store visible text in shapes, text boxes, headers,
    # footers, or content structures that are not exposed as normal paragraphs.
    # Falling back to raw Word XML makes uploads more forgiving before the note
    # is sent into the later API-ready fact extraction stage.
    return extract_docx_text_from_xml(file_bytes)


def read_uploaded_note(uploaded_file) -> str:
    """Convert a supported uploaded note file into plain text.

    Args:
        uploaded_file: Streamlit uploaded file object for a `.txt`, `.md`, or
        `.docx` clinical note.

    Returns:
        Plain text extracted from the uploaded note. Returns an empty string if
        the file type is unsupported, the document cannot be parsed, or no text
        can be extracted.

    CLARITY pipeline role:
        Normalizes multiple note input formats into one raw-text representation
        for Step 1. This keeps downstream API-ready stages simpler: future fact
        extraction services can receive note text without needing to know
        whether the user uploaded text, markdown, or Word.
    """
    suffix = Path(uploaded_file.name).suffix.lower()
    file_bytes = uploaded_file.getvalue()

    if suffix in {".txt", ".md"}:
        return file_bytes.decode("utf-8", errors="replace")

    if suffix == ".docx":
        try:
            note_text = extract_docx_text(file_bytes)
        except Exception as exc:  # python-docx can raise several parser errors.
            st.error(f"Unable to parse Word document `{uploaded_file.name}`: {exc}")
            return ""

        if not note_text:
            st.warning(
                "No readable text was found in Word document "
                f"`{uploaded_file.name}`. If this file is a scanned image or "
                "image-only export, convert it with OCR before uploading."
            )
        return note_text

    st.warning(
        "Unsupported file type. Please upload a `.txt`, `.md`, or `.docx` note."
    )
    return ""


def handle_case_import(sample_path: Path) -> tuple[str, str, str]:
    """Render the clinical note import controls and return the loaded case.

    Args:
        sample_path: Path to the built-in sample note used by the demo flow.

    Returns:
        A tuple of `(case_status, note_text, pipeline_mode)`. `case_status` is
        one of "No case loaded", "Sample case loaded", or
        "Uploaded case loaded". `note_text` contains the loaded raw clinical
        note text, or an empty string when no case is loaded. `pipeline_mode`
        is "demo", "api", or "none".

    CLARITY pipeline role:
        Implements Step 1, where a de-identified clinical note enters the
        patient explainer workflow. The prototype accepts text, markdown, and
        Word documents, then normalizes them to plain text so structured fact
        extraction remains a separate API-ready step.
    """
    render_step_header(
        1,
        "Import Clinical Note",
        "Load a de-identified note from the sample case or from a local text, markdown, or Word file.",
    )

    sample_col, upload_col = st.columns([1, 1], gap="large")
    with sample_col:
        st.markdown("**Built-in demo case**")
        st.caption("Best for a fast interview walkthrough.")
        use_sample = st.checkbox("Use sample de-identified case note", value=True)

    with upload_col:
        st.markdown("**Upload local note**")
        st.caption("Supports `.txt`, `.md`, and `.docx` files for this prototype.")
        uploaded_file = st.file_uploader(
            "Upload a de-identified clinical note",
            type=["txt", "md", "docx"],
            help="Current prototype supports plain text, markdown, and Word documents.",
            label_visibility="collapsed",
        )

    if uploaded_file is not None:
        note_text = read_uploaded_note(uploaded_file)
        if note_text:
            st.success(f"Uploaded case loaded: {uploaded_file.name}")
            st.info(
                "Uploaded case loaded in API mode. Steps 2-4 will use "
                "configured API endpoints instead of local demo assets."
            )
            return "Uploaded case loaded", note_text, "api"
        st.warning("Uploaded case could not be loaded.")
        return "No case loaded", "", "none"

    if use_sample:
        note_text = load_sample_note(sample_path)
        if note_text:
            st.success("Sample case loaded")
            st.caption("Sample case uses local demo assets.")
            return "Sample case loaded", note_text, "demo"

    st.info("No case loaded")
    return "No case loaded", "", "none"


def display_safety_notice() -> None:
    """Display the prototype safety notice.

    Args:
        None.

    Returns:
        None. The notice is rendered directly into the Streamlit page.

    CLARITY pipeline role:
        Reinforces that the dashboard is a patient education prototype, not a
        source of medical advice. This notice frames every downstream step:
        fact review, script presentation, and cached video display should all be
        understood as materials requiring clinician review before real use.
    """
    with st.expander("Prototype safety notice", expanded=False):
        st.warning(
            "This dashboard is for demonstration and patient education design "
            "only. It is not medical advice, and final content should be "
            "reviewed by clinicians before patient-facing use."
        )


def display_demo_overview() -> None:
    """Display a concise overview of the step-by-step demo flow.

    Args:
        None.

    Returns:
        None. The overview is rendered directly into the Streamlit page.

    CLARITY pipeline role:
        Helps evaluators understand the prototype within the first few seconds
        of the demo. The overview anchors the app as a workflow demonstration
        rather than a standalone video player or generic dashboard.
    """
    st.markdown(
        """
        <div class="clarity-hero">
            <div class="clarity-hero-title">Patient-facing explainer workflow</div>
            <p class="clarity-hero-text">
                Import a de-identified note, review the shared fact base, choose
                one of five communication modes, then inspect the selected script,
                cached video slot, transparency notes, and limitations. The current
                demo uses local cached assets, but Steps 1-4 are structured so they
                can later be connected to API-backed extraction, generation,
                verification, and video services.
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def display_case_snapshot(fact_base: dict, pipeline_mode: str) -> None:
    """Render a patient-facing snapshot from the shared structured fact base.

    Args:
        fact_base: Dictionary loaded from `data/extracted_fact_base.json` in
        demo mode, or returned by the fact extraction API in uploaded-case mode.
        pipeline_mode: "demo", "api", or "none".

    Returns:
        None. The snapshot is rendered directly into the Streamlit page.

    CLARITY pipeline role:
        Implements Step 2 by showing the structured facts that will anchor all
        explanation versions. The five communication modes should share this
        same fact base so tone, language, and uncertainty framing can vary
        without allowing facts to drift between versions.
    """
    render_step_header(
        2,
        "Review Shared Fact Base",
        "Inspect the clinician-verifiable facts that anchor every explanation version.",
    )
    if pipeline_mode == "api":
        st.info(
            "Uploaded case API mode: the structured fact base should come from "
            "`FACT_EXTRACTION_API_URL`. Local sample facts are not used."
        )
    else:
        st.info(
            "For this prototype, the structured fact base is loaded from a "
            "pre-verified JSON file. In a production workflow, this step would "
            "be generated from the clinical note and reviewed by a clinician."
        )

    if not fact_base:
        st.warning("No structured fact base is available yet.")
        return

    # Fact extraction and style-controlled script presentation are separated so
    # every version can preserve the same clinical content while changing only
    # language, tone, complexity, and uncertainty framing.
    st.markdown(
        "**Shared fact base:** all explanation versions should be generated "
        "from this same structured source to reduce factual drift."
    )

    diagnosis_col, stage_col, age_col = st.columns(3)
    with diagnosis_col:
        render_fact_card("Diagnosis", fact_base.get("diagnosis", "Not available"))
    with stage_col:
        render_fact_card("Stage", fact_base.get("stage", "Not available"))
    with age_col:
        render_fact_card("Patient age", fact_base.get("patient_age", "Not available"))

    area_col, treatment_col = st.columns(2)
    with area_col:
        st.markdown("**Main areas involved**")
        st.markdown('<div class="compact-list">', unsafe_allow_html=True)
        for area in fact_base.get("main_areas_involved", []):
            st.markdown(f"- {area}")
        st.markdown("</div>", unsafe_allow_html=True)

    with treatment_col:
        render_fact_card(
            "Current treatment",
            fact_base.get("current_treatment", "Not available"),
        )
        render_fact_card(
            "Recommended next treatment",
            fact_base.get("recommended_next_treatment", "Not available"),
        )

    st.markdown("**Key uncertainty / monitoring points**")
    st.markdown('<div class="compact-list">', unsafe_allow_html=True)
    for point in fact_base.get("uncertainty_points", []):
        st.markdown(f"- {point}")
    st.markdown("</div>", unsafe_allow_html=True)


def display_key_takeaways(fact_base: dict) -> None:
    """Display concise patient-facing takeaways from the shared fact base.

    Args:
        fact_base: Dictionary loaded from `data/extracted_fact_base.json`.

    Returns:
        None. The takeaways are rendered directly into the Streamlit page.

    CLARITY pipeline role:
        Provides a short patient-facing summary that is anchored to the same
        structured facts used by every explanation mode. This gives evaluators
        a stable reference point before they compare style-controlled scripts.
    """
    st.markdown("**Key Takeaways**")
    takeaways = fact_base.get("key_takeaways", [])
    if not takeaways:
        st.warning("No key takeaways are available in the shared fact base.")
        return

    for takeaway in takeaways:
        st.markdown(f"- {takeaway}")


def display_questions_for_care_team(fact_base: dict) -> None:
    """Display suggested questions for the patient to ask their care team.

    Args:
        fact_base: Dictionary loaded from `data/extracted_fact_base.json`.

    Returns:
        None. The questions are rendered directly into the Streamlit page.

    CLARITY pipeline role:
        Turns the verified fact base into practical discussion prompts. The
        questions are educational and preparatory, not direct medical advice,
        which helps keep the prototype aligned with its patient education role.
    """
    st.markdown("**Questions for Your Care Team**")
    questions = fact_base.get("questions_for_care_team", [])
    if not questions:
        st.warning("No care team questions are available in the shared fact base.")
        return

    for question in questions:
        st.markdown(f"- {question}")


def display_version_metadata(metadata: dict) -> None:
    """Render the selected explanation mode metadata.

    Args:
        metadata: Dictionary describing the currently selected explanation
        version, including title, language, communication goal, tone, technical
        detail, and uncertainty detail.

    Returns:
        None. The version metadata is rendered directly into the Streamlit page.

    CLARITY pipeline role:
        Implements Step 3 by making the style-control layer visible. The
        selected mode changes how an explanation is presented, but it should not
        change the underlying shared clinical facts loaded in Step 2.
    """
    render_step_header(
        3,
        "Choose Explanation Mode",
        "Compare how the same facts can be presented with different communication goals.",
    )

    if not metadata:
        st.warning("No explanation mode metadata is available yet.")
        return

    st.markdown(
        f"""
        <div class="mode-summary">
            <div class="mode-title">{escape(metadata.get('title', 'Selected version'))}</div>
            <div class="muted-note">{escape(metadata.get('communication_goal', 'No communication goal provided.'))}</div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    language_col, tone_col = st.columns(2)
    language_col.markdown(f"**Language:** {metadata.get('language', 'Not available')}")
    tone_col.markdown(f"**Tone:** {metadata.get('tone', 'Not available')}")

    detail_col, uncertainty_col = st.columns(2)
    detail_col.markdown(
        f"**Technical detail:** {metadata.get('technical_detail', 'Not available')}"
    )
    uncertainty_col.markdown(
        f"**Uncertainty detail:** {metadata.get('uncertainty_detail', 'Not available')}"
    )


def select_explanation_mode(version_metadata: dict) -> tuple[str, dict]:
    """Render the sidebar explanation mode selector.

    Args:
        version_metadata: Dictionary loaded from `data/version_metadata.json`,
        keyed by stable mode identifiers.

    Returns:
        A tuple of `(selected_mode_key, selected_mode_metadata)`. If metadata is
        unavailable, returns an empty key and empty dictionary.

    CLARITY pipeline role:
        Allows the evaluator to switch between the five predefined presentation
        styles required by the interview task. This function is intentionally
        limited to selecting cached/predefined versions; later production work
        could connect these choices to real-time generation after verification
        safeguards are in place.
    """
    st.sidebar.header("Explanation mode")

    if not version_metadata:
        st.sidebar.warning("Version metadata is unavailable.")
        return "", {}

    mode_keys = list(version_metadata.keys())
    labels = {
        key: version_metadata[key].get("title", key.replace("_", " ").title())
        for key in mode_keys
    }
    selected_key = st.sidebar.selectbox(
        "Choose a predefined explanation mode",
        options=mode_keys,
        format_func=lambda key: labels[key],
    )
    return selected_key, version_metadata[selected_key]


def display_future_personalization_controls() -> dict:
    """Render placeholder controls for future patient preference selection.

    Args:
        None.

    Returns:
        A dictionary mapping stable preference names to boolean selections.

    CLARITY pipeline role:
        Shows how participants might eventually choose presentation preferences
        such as simpler language, more technical detail, clearer uncertainty, a
        more reassuring tone, or Spanish language. In this prototype the
        controls intentionally do not regenerate content; evaluators should use
        the five predefined modes until a verified real-time generation pipeline
        exists.
    """
    st.sidebar.header("Future personalization controls")
    st.sidebar.info(
        "These controls are placeholders for future real-time generation. In "
        "this prototype, please use the five predefined modes below."
    )

    return {
        "Make it easier to understand": st.sidebar.checkbox(
            "Make it easier to understand"
        ),
        "Include more technical detail": st.sidebar.checkbox(
            "Include more technical detail"
        ),
        "Explain uncertainty more clearly": st.sidebar.checkbox(
            "Explain uncertainty more clearly"
        ),
        "Use a more reassuring tone": st.sidebar.checkbox(
            "Use a more reassuring tone"
        ),
        "Spanish language": st.sidebar.checkbox("Spanish language"),
    }


def display_personalization_summary(preferences: dict) -> None:
    """Show a non-generative summary of selected future preferences.

    Args:
        preferences: Dictionary returned by
        `display_future_personalization_controls`.

    Returns:
        None. The summary is rendered into the Streamlit page.

    CLARITY pipeline role:
        Makes placeholder personalization choices visible without changing the
        current transcript or video. This avoids implying that the prototype is
        performing live AI generation or clinical verification.
    """
    selected_preferences = [
        label for label, is_selected in preferences.items() if is_selected
    ]

    st.markdown("**Selected preferences summary**")
    if selected_preferences:
        for preference in selected_preferences:
            st.markdown(f"- {preference}")
    else:
        st.markdown(
            '<p class="muted-note">No future personalization preferences selected.</p>',
            unsafe_allow_html=True,
        )

    st.caption(
        "Preference controls are placeholders only; the selected predefined "
        "mode determines the script and cached video shown in this prototype."
    )


def display_script(script_path: Path) -> None:
    """Display the selected explanation script or a graceful missing-file note.

    Args:
        script_path: Path to the markdown transcript associated with the
        selected explanation mode.

    Returns:
        None. The transcript or fallback warning is rendered into the page.

    CLARITY pipeline role:
        Implements the transcript half of Step 4. In the prototype, scripts are
        loaded from cached markdown files rather than generated live. This keeps
        the interview demo reproducible and reinforces that style-controlled
        explanation should happen after fact extraction and verification.
    """
    with st.expander("View transcript / script", expanded=False):
        script_text = load_markdown(script_path)
        if script_text:
            st.markdown(script_text)
        else:
            st.warning(
                "No transcript is available for this selected version yet. "
                "Add a markdown file at the path listed in version metadata."
            )


def display_script_text(script_text: str, source_label: str) -> None:
    """Display API-generated script text in the transcript expander.

    Args:
        script_text: Generated script markdown or plain text returned by the
        script generation API.
        source_label: Short label describing where the script came from.

    Returns:
        None. The script text or fallback warning is rendered into the page.

    CLARITY pipeline role:
        Supports uploaded-case API mode for Step 4. It keeps generated scripts
        separate from local demo markdown files so the dashboard does not show
        sample-case content for user-uploaded clinical notes.
    """
    with st.expander("View transcript / script", expanded=True):
        st.caption(source_label)
        if script_text:
            st.markdown(script_text)
        else:
            st.warning("No API-generated script is available yet.")


def display_video_or_placeholder(video_path: Path) -> None:
    """Display a cached explainer video or an intentional placeholder panel.

    Args:
        video_path: Path to the cached video file for the selected explanation
        mode.

    Returns:
        None. The video or placeholder message is rendered into the page.

    CLARITY pipeline role:
        Implements the video half of Step 4. The prototype intentionally reads
        cached videos rather than generating videos in real time because video
        generation is slow, expensive, and should follow script verification.
        This function makes that design choice visible while keeping the
        dashboard ready to host generated videos once they are available.
    """
    st.markdown("**Cached video explainer**")

    if video_path.exists():
        st.video(str(video_path))
        return

    # Missing cached video files are expected at this prototype stage. The
    # fallback should look deliberate so evaluators understand the dashboard is
    # prepared for videos, not broken by their absence.
    st.markdown(
        """
        <div class="video-placeholder">
            <div class="video-placeholder-title">Video not generated yet</div>
            <p class="video-placeholder-text">
                This space will display the 5-10 minute NotebookLM-style
                explainer video for the selected version. For the current
                prototype, the dashboard is prepared to host the generated video
                once available.
            </p>
        </div>
        """,
        unsafe_allow_html=True,
    )


def display_api_video_response(video_response: dict) -> None:
    """Display the result of the video API call for uploaded-case mode.

    Args:
        video_response: JSON object returned by the configured video generation
        or retrieval API.

    Returns:
        None. A playable video, status panel, or API response summary is
        rendered into the Streamlit page.

    CLARITY pipeline role:
        Supports Step 4 API mode. Uploaded cases should use API-provided video
        outputs rather than local sample cached videos.
    """
    st.markdown("**API video output**")
    if not video_response:
        st.warning("No video API response is available yet.")
        return

    video_url = extract_response_text(video_response, ["video_url", "url"])
    if video_url:
        st.video(video_url)
        return

    video_file = extract_response_text(video_response, ["video_file", "video_path"])
    if video_file:
        candidate_path = Path(video_file)
        if not candidate_path.is_absolute():
            candidate_path = APP_ROOT / candidate_path
        if candidate_path.exists():
            st.video(str(candidate_path))
            return
        st.warning(f"Video API returned a file path that does not exist: {video_file}")

    status = extract_response_text(video_response, ["status", "message"])
    if status:
        st.info(status)

    with st.expander("Raw video API response", expanded=False):
        st.json(video_response)


def display_pipeline_transparency() -> None:
    """Display the prototype generation workflow and transparency rationale.

    Args:
        None.

    Returns:
        None. The pipeline explanation is rendered inside a Streamlit expander.

    CLARITY pipeline role:
        Implements the transparency portion of Step 5. It helps interview
        evaluators understand that CLARITY separates note import, fact-base
        creation, style-controlled script drafting, verification, and video
        generation instead of treating the final explainer as a black box.
    """
    with st.expander("How this prototype was generated", expanded=False):
        st.markdown(
            """
1. Import or select a de-identified clinical note.
2. Load a structured clinical fact base.
3. Convert facts into a patient-facing outline.
4. Generate five style-controlled scripts from the same fact base.
5. Verify scripts against the source note for unsupported claims or missing critical information.
6. Convert verified scripts into video explainers.
7. Present cached videos and transcripts through this dashboard.
"""
        )
        st.info(
            "All five versions should come from one shared fact base. This "
            "reduces factual drift when the communication style changes across "
            "language, trust-building, technical detail, and uncertainty framing."
        )
        st.markdown(
            """
**Future API handoff:** Step 1 can send the imported note to a secure ingestion
service, Step 2 can call fact extraction plus clinician review tooling, Step 3
can call style-controlled script generation and verification, and Step 4 can
retrieve cached verified video outputs after generation is complete.
"""
        )


def display_demo_guide() -> None:
    """Display a short guide for interview evaluators.

    Args:
        None.

    Returns:
        None. The guide is rendered in a Streamlit expander.

    CLARITY pipeline role:
        Provides a compact walkthrough that an evaluator can follow during a
        live or recorded demo. It highlights the intended interactions without
        adding visible instructional clutter to the main patient-facing flow.
    """
    with st.expander("Demo Guide for Evaluators", expanded=False):
        st.markdown(
            """
- Load the sample clinical note.
- Review the shared clinical fact base.
- Switch between five explanation modes.
- Open the transcript panel.
- Notice that the video panel is ready for cached generated videos.
- Review limitations and next steps.
"""
        )


def display_limitations(writeup_path: Path) -> None:
    """Display limitations and next steps from a maintained markdown writeup.

    Args:
        writeup_path: Path to the limitations and next steps markdown file.

    Returns:
        None. The writeup is rendered inside a Streamlit expander.

    CLARITY pipeline role:
        Completes the limitations portion of Step 5 by making prototype
        boundaries explicit. This is important for an interview demo because it
        distinguishes the current local dashboard from a clinically deployable
        production system that would require verification, audit trails,
        privacy controls, translation review, and comprehension testing.
    """
    with st.expander("Limitations and next steps", expanded=False):
        writeup = load_markdown(writeup_path)
        if writeup:
            st.markdown(writeup)
        else:
            st.warning(
                "Limitations writeup is not available. Add the markdown file "
                "under `writeup/` so evaluators can review prototype boundaries."
            )


def main() -> None:
    """Render the CLARITY Streamlit application shell.

    This function configures the page and coordinates the visible workflow. At
    this milestone it renders the clinical note import step; later milestones
    will add shared fact-base review, explanation mode selection, script/video
    review, and transparency/limitations documentation.

    Args:
        None.

    Returns:
        None. Streamlit renders the page as a side effect.

    CLARITY pipeline role:
        Provides the stable application entrypoint that future milestones will
        use to present the patient-facing explainer workflow.
    """
    st.set_page_config(
        page_title="CLARITY Patient Explainer Dashboard",
        layout="wide",
    )
    inject_custom_css()
    api_config = load_environment_config(ENV_PATH)

    st.title("CLARITY Patient Explainer Dashboard")
    st.caption(
        "A prototype for turning clinician-verified notes into patient-facing "
        "video explanations."
    )

    st.markdown(
        "This local demo walks through the CLARITY pipeline one step at a time, "
        "from note import to patient-facing explanation materials."
    )

    display_demo_overview()
    render_workflow_strip()
    display_safety_notice()

    st.divider()
    case_status, note_text, pipeline_mode = handle_case_import(SAMPLE_NOTE_PATH)

    st.markdown(f"**Current case status:** {case_status}")
    if pipeline_mode == "api":
        st.markdown("**Pipeline mode:** API mode for uploaded case")
    elif pipeline_mode == "demo":
        st.markdown("**Pipeline mode:** Local demo mode")
    if note_text:
        with st.expander("Raw clinical note preview", expanded=False):
            st.text_area(
                "Imported note text",
                value=note_text,
                height=260,
                disabled=True,
                label_visibility="collapsed",
            )

    st.divider()
    version_metadata = load_json(VERSION_METADATA_PATH)
    selected_mode_key, selected_mode_metadata = select_explanation_mode(
        version_metadata
    )
    personalization_preferences = display_future_personalization_controls()
    display_api_config_status(api_config)

    note_hash = get_note_hash(note_text) if note_text else ""
    if pipeline_mode == "api" and note_hash:
        ingestion_cache_key = (
            f"api_ingestion:{note_hash}:{api_config.get('NOTE_INGESTION_API_URL', '')}"
        )
        if ingestion_cache_key not in st.session_state:
            with st.spinner("Calling note ingestion API..."):
                st.session_state[ingestion_cache_key] = call_note_ingestion_api(
                    note_text,
                    api_config,
                )
        if st.session_state[ingestion_cache_key]:
            with st.expander("Step 1 API ingestion response", expanded=False):
                st.json(st.session_state[ingestion_cache_key])

    if pipeline_mode == "demo":
        fact_base = load_json(FACT_BASE_PATH)
    elif pipeline_mode == "api" and note_hash:
        fact_cache_key = (
            f"api_fact_base:{note_hash}:{api_config.get('FACT_EXTRACTION_API_URL', '')}"
        )
        if fact_cache_key not in st.session_state:
            with st.spinner("Calling fact extraction API..."):
                st.session_state[fact_cache_key] = call_fact_extraction_api(
                    note_text,
                    api_config,
                )
        fact_base = st.session_state[fact_cache_key]
    else:
        fact_base = {}

    display_case_snapshot(fact_base, pipeline_mode)
    if fact_base:
        takeaway_col, question_col = st.columns(2)
        with takeaway_col:
            display_key_takeaways(fact_base)
        with question_col:
            display_questions_for_care_team(fact_base)

    st.divider()
    display_version_metadata(selected_mode_metadata)
    display_personalization_summary(personalization_preferences)

    st.divider()
    render_step_header(
        4,
        "Review Script and Cached Video",
        "Open the selected transcript and confirm the video slot is ready for cached generated media.",
    )
    script_file = selected_mode_metadata.get("script_file") if selected_mode_metadata else ""
    video_file = selected_mode_metadata.get("video_file") if selected_mode_metadata else ""
    if pipeline_mode == "api":
        if not fact_base:
            st.warning(
                "API mode is active, but no fact base is available. Configure "
                "and run `FACT_EXTRACTION_API_URL` before generating scripts."
            )
        else:
            preferences_key = json.dumps(
                personalization_preferences,
                sort_keys=True,
            )
            script_cache_key = (
                f"api_script:{note_hash}:{selected_mode_key}:"
                f"{api_config.get('SCRIPT_GENERATION_API_URL', '')}:{preferences_key}"
            )
            if script_cache_key not in st.session_state:
                with st.spinner("Calling script generation API..."):
                    st.session_state[script_cache_key] = call_script_generation_api(
                        note_text,
                        fact_base,
                        selected_mode_key,
                        selected_mode_metadata,
                        personalization_preferences,
                        api_config,
                    )
            script_text = st.session_state[script_cache_key]
            display_script_text(script_text, "Source: script generation API")

            verification_report = {}
            if script_text:
                script_hash = sha256(script_text.encode("utf-8")).hexdigest()[:16]
                verification_cache_key = (
                    f"api_verification:{note_hash}:{selected_mode_key}:"
                    f"{api_config.get('VERIFICATION_API_URL', '')}:{script_hash}"
                )
                if verification_cache_key not in st.session_state:
                    with st.spinner("Calling verification API..."):
                        st.session_state[verification_cache_key] = (
                            call_verification_api(
                                note_text,
                                fact_base,
                                script_text,
                                selected_mode_key,
                                api_config,
                            )
                        )
                verification_report = st.session_state[verification_cache_key]
                if verification_report:
                    with st.expander("Verification API response", expanded=False):
                        st.json(verification_report)

                video_cache_key = (
                    f"api_video:{note_hash}:{selected_mode_key}:"
                    f"{api_config.get('VIDEO_GENERATION_API_URL', '')}:{script_hash}"
                )
                if video_cache_key not in st.session_state:
                    with st.spinner("Calling video generation API..."):
                        st.session_state[video_cache_key] = call_video_generation_api(
                            fact_base,
                            script_text,
                            selected_mode_key,
                            selected_mode_metadata,
                            verification_report,
                            api_config,
                        )
                display_api_video_response(st.session_state[video_cache_key])
            else:
                st.warning("Skipping video API call because no script was generated.")
    else:
        if script_file:
            display_script(APP_ROOT / script_file)
        else:
            st.warning("No script file is configured for the selected mode.")

        if video_file:
            display_video_or_placeholder(APP_ROOT / video_file)
        else:
            st.warning("No video file is configured for the selected mode.")

    st.divider()
    render_step_header(
        5,
        "Review Pipeline Transparency, Limitations, and Next Steps",
        "Show how the prototype was generated and what must happen before real patient-facing use.",
    )
    display_demo_guide()
    display_pipeline_transparency()
    display_limitations(LIMITATIONS_PATH)


if __name__ == "__main__":
    main()
