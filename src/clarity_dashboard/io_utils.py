"""File and document loading helpers for the CLARITY dashboard prototype."""

import json
import zipfile
from io import BytesIO
from pathlib import Path
from xml.etree import ElementTree

import streamlit as st
from docx import Document


def load_json(path: Path) -> dict:
    """Load a JSON object from disk with Streamlit-friendly error handling.

    Args:
        path: Path to a JSON file expected to contain a top-level object.

    Returns:
        A dictionary parsed from the JSON file. Returns an empty dictionary when
        the file is missing, unreadable, malformed, or does not contain a JSON
        object.

    CLARITY pipeline role:
        Provides a safe loader for structured prototype assets such as the
        clinician-verifiable fact base and version metadata. Safe loading
        matters because this dashboard is meant to stay demoable while scripts,
        prompts, and videos are being iterated.
    """
    try:
        data = json.loads(path.read_text(encoding="utf-8"))
    except FileNotFoundError:
        # A missing data file should be visible to the evaluator without
        # collapsing the rest of the prototype page.
        st.warning(f"Required JSON file was not found: {path}")
        return {}
    except json.JSONDecodeError as exc:
        st.error(f"Invalid JSON in {path.name}: {exc}")
        return {}
    except OSError as exc:
        st.error(f"Unable to read JSON file {path.name}: {exc}")
        return {}

    if not isinstance(data, dict):
        st.error(f"Expected {path.name} to contain a JSON object.")
        return {}
    return data


def load_markdown(path: Path) -> str:
    """Load a markdown document from disk with graceful error handling.

    Args:
        path: Path to the markdown file that should be rendered in the app.

    Returns:
        The markdown text when it can be read, or an empty string when the file
        is missing or unreadable.

    CLARITY pipeline role:
        Supports transcript, prompt, and writeup display throughout the
        prototype. Markdown assets are intentionally stored as editable files so
        scripts and documentation can be revised without changing Streamlit
        application logic.
    """
    try:
        return path.read_text(encoding="utf-8")
    except FileNotFoundError:
        # Missing files are expected during iteration; the dashboard should
        # point to the missing asset instead of failing mid-demo.
        st.warning(f"Markdown file was not found: {path}")
    except OSError as exc:
        st.error(f"Unable to read markdown file {path.name}: {exc}")
    return ""


def load_sample_note(path: Path) -> str:
    """Load the built-in de-identified sample clinical note.

    Args:
        path: Path to the sample clinical note text file.

    Returns:
        The note text when the file can be read, or an empty string when the
        file is missing or unreadable.

    CLARITY pipeline role:
        Supports Step 1 of the prototype workflow by giving interview
        evaluators a reliable built-in case to load before they test upload
        behavior. This keeps the demo runnable even when the evaluator does not
        have a local clinical note prepared.
    """
    try:
        return path.read_text(encoding="utf-8")
    except FileNotFoundError:
        # Missing demo assets should not crash the app; graceful fallback keeps
        # the prototype usable while clearly showing what needs to be restored.
        st.warning(f"Sample note file was not found: {path}")
    except OSError as exc:
        st.error(f"Unable to read sample note: {exc}")
    return ""


def extract_docx_text_from_xml(file_bytes: bytes) -> str:
    """Extract Word text nodes directly from the `.docx` XML package.

    Args:
        file_bytes: Raw bytes from the uploaded Word document.

    Returns:
        Plain text assembled from WordprocessingML text nodes across document
        body, headers, footers, footnotes, endnotes, and other Word XML parts.
        Returns an empty string when no text nodes are present.

    CLARITY pipeline role:
        Provides a fallback parser for Step 1 when clinical note text is stored
        in Word structures that `python-docx` does not expose through ordinary
        paragraphs or tables. This keeps the prototype robust for real demo
        files while still normalizing everything to plain text before
        API-backed fact extraction.
    """
    text_blocks: list[str] = []
    namespace = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

    with zipfile.ZipFile(BytesIO(file_bytes)) as archive:
        xml_part_names = [
            name
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        ]

        for part_name in xml_part_names:
            try:
                root = ElementTree.fromstring(archive.read(part_name))
            except ElementTree.ParseError:
                continue

            text_values = [
                node.text
                for node in root.findall(".//w:t", namespace)
                if node.text and node.text.strip()
            ]
            if text_values:
                text_blocks.append(" ".join(value.strip() for value in text_values))

    return "\n\n".join(text_blocks)


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract readable text from an uploaded Word `.docx` document.

    Args:
        file_bytes: Raw bytes from the uploaded Word document.

    Returns:
        Plain text extracted from document paragraphs, table cells, and a
        WordprocessingML XML fallback. Returns an empty string when the document
        has no readable text nodes.

    CLARITY pipeline role:
        Extends Step 1 so clinical notes can enter the prototype as Word
        documents, which is a more realistic format for interview demos and
        clinical handoffs. The function still converts the upload into plain
        text before later steps, keeping fact extraction separate from document
        format parsing.
    """
    document = Document(BytesIO(file_bytes))
    text_blocks: list[str] = []

    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()
        if paragraph_text:
            text_blocks.append(paragraph_text)

    for table in document.tables:
        for row in table.rows:
            cell_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if cell_text:
                text_blocks.append(cell_text)

    paragraph_table_text = "\n\n".join(text_blocks)
    if paragraph_table_text:
        return paragraph_table_text

    # Some Word files store visible text in shapes, text boxes, headers,
    # footers, or content structures that are not exposed as normal paragraphs.
    # Falling back to raw Word XML makes uploads more forgiving before the note
    # is sent into the later API-ready fact extraction stage.
    return extract_docx_text_from_xml(file_bytes)


def read_uploaded_note(uploaded_file) -> str:
    """Convert a supported uploaded note file into plain text.

    Args:
        uploaded_file: Streamlit uploaded file object for a `.txt`, `.md`, or
        `.docx` clinical note.

    Returns:
        Plain text extracted from the uploaded note. Returns an empty string if
        the file type is unsupported, the document cannot be parsed, or no text
        can be extracted.

    CLARITY pipeline role:
        Normalizes multiple note input formats into one raw-text representation
        for Step 1. This keeps downstream API-ready stages simpler: future fact
        extraction services can receive note text without needing to know
        whether the user uploaded text, markdown, or Word.
    """
    suffix = Path(uploaded_file.name).suffix.lower()
    file_bytes = uploaded_file.getvalue()

    if suffix in {".txt", ".md"}:
        return file_bytes.decode("utf-8", errors="replace")

    if suffix == ".docx":
        try:
            note_text = extract_docx_text(file_bytes)
        except Exception as exc:  # python-docx can raise several parser errors.
            st.error(f"Unable to parse Word document `{uploaded_file.name}`: {exc}")
            return ""

        if not note_text:
            st.warning(
                "No readable text was found in Word document "
                f"`{uploaded_file.name}`. If this file is a scanned image or "
                "image-only export, convert it with OCR before uploading."
            )
        return note_text

    st.warning(
        "Unsupported file type. Please upload a `.txt`, `.md`, or `.docx` note."
    )
    return ""
